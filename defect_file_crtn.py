#######################################################################################
# Author : Utkarsh Singh
# Project : Daily Utility
# Module : Window to take Defect Details and add it to word and make defect files
#######################################################################################
from tkinter import *
from tkinter import Text
from tkinter import ttk
from tkinter import messagebox
import os
import docx
from docx.shared import Cm
from docx import Document
import Utility_Launch
import datetime
from Logging_files import defect_info


class defect_file(Toplevel):
    def __init__(self):
        Toplevel.__init__(self)
        self.geometry("550x920+700+90")
        self.title("Defect Details !")
        self.resizable(False, False)
        # Top hold i and Name Label
        self.top_defect_frame = Frame(self, height=120, bg='white').pack(fill='both')
        try:
            self.top_defect_image = PhotoImage(file=os.path.join(Utility_Launch.path, 'Image_01/defect.png'))
            self.label_defect_photo = ttk.Label(self, image=self.top_defect_image, background="white")
            self.label_defect_photo.place(x=30, y=30)
        except:
            pass
        self.defect_label = ttk.Label(self, text="Caught It ", font=("Times New Roman", 15)
                                      , background="white").place(x=200, y=40)
        self.defect_label2 = ttk.Label(self, text="Raise It!!!", font=("Times New Roman", 10)
                                       , background="white").place(x=200, y=80)

        # Bottom frame to hold buttons and entry texts
        self.btm_defect_frame = Frame(self, height=900, bg='#50946a').pack(fill='both')
        self.defect_label2 = ttk.Label(self, text="Defect Details", font=("Times New Roman", 16)
                                       , background="#50946a").place(x=190, y=140)
        self.defect_label3 = ttk.Label(self, text="Enter G-ID : ", font=("Times New Roman", 12)
                                       , background="#50946a").place(x=10, y=180)
        self.id_entry = ttk.Entry(self, width=15)
        self.id_entry.insert(0, "G")
        self.id_entry.place(x=180, y=180)
        self.id_entry.focus_force()

        self.defect_timestamp = ttk.Label(self, text="Enter Time in CET", font=("Times New Roman", 12)
                                          , background="#50946a").place(x=10, y=220)
        self.defect_timestamp_entry = ttk.Entry(self, width=15)
        self.defect_timestamp_entry.place(x=180, y=220)
        self.defect_label3 = ttk.Label(self, text="Enter App. Name : ", font=("Times New Roman", 12)
                                       , background="#50946a").place(x=10, y=260)
        self.app_name_entry = ttk.Entry(self, width=30)
        self.app_name_entry.place(x=180, y=260)

        self.defect_label4 = ttk.Label(self, text="Enter Actual : ", font=("Times New Roman", 12)
                                       , background="#50946a").place(x=10, y=320)
        self.actual_result = Text(self, height=8, width=23, padx=5, pady=4, relief=RAISED, border=4,
                                  wrap=WORD)
        self.actual_result.place(x=180, y=320)

        self.defect_label5 = ttk.Label(self, text="Enter Expected : ", font=("Times New Roman", 12)
                                       , background="#50946a").place(x=10, y=510)
        self.expected_result = Text(self, height=8, width=23, padx=5, pady=4, relief=RAISED, border=4,
                                    wrap=WORD)
        self.expected_result.place(x=180, y=510)

        self.defect_label6 = ttk.Label(self, text="Enter File Type : ", font=("Times New Roman", 12)
                                       , background="#50946a").place(x=10, y=700)
        self.file_name_entry = ttk.Entry(self, width=30)
        self.file_name_entry.place(x=180, y=700)

        self.defect_label7 = ttk.Label(self, text="Enter File Number : ", font=("Times New Roman", 12)
                                       , background="#50946a").place(x=10, y=760)
        self.file_number_entry = ttk.Entry(self, width=15)
        self.file_number_entry.place(x=180, y=760)

        self.defect_raise = ttk.Button(self, width=14, text='DONE', command=self.write_in_defect_file)
        self.defect_raise.place(x=120, y=820)
        self.close_defect_raise = ttk.Button(self, width=15, text='    Close    ', command=self.close_wndw)
        self.close_defect_raise.place(x=255, y=820)

    def close_wndw(self):
        self.destroy()

    def defect_template(self):

        doc_tem = Document()
        doc_para0 = doc_tem.add_paragraph()
        doc_para0.add_run("Execution Info ").bold = True

        # Table for holding User Id and App name
        info_table = doc_tem.add_table(rows=3, cols=2)
        info_table.style = "Table Grid"
        frst_cell = info_table.rows[0].cells
        frst_cell[0].text = 'User Id'
        frst_cell[1].text = self.id_entry.get()

        scnd_cell = info_table.rows[1].cells
        scnd_cell[0].text = 'Time Stamp'
        current = str(datetime.datetime.now())
        date, times = current.split()
        time_stamp = date + " " + self.defect_timestamp_entry.get()
        scnd_cell[1].text = time_stamp

        third_cell = info_table.rows[2].cells
        third_cell[0].text = 'Application Name'
        third_cell[1].text = self.app_name_entry.get()

        # Actual and Expected Section
        doc_para = doc_tem.add_paragraph()
        doc_para.add_run("Actual Result").bold = True
        doc_para_des = doc_tem.add_paragraph()
        doc_para_des.add_run(self.actual_result.get('1.0', END))

        doc_para2 = doc_tem.add_paragraph()
        doc_para2.add_run("Expected Result").bold = True
        doc_para_des2 = doc_tem.add_paragraph()
        doc_para_des2.add_run(self.expected_result.get('1.0', END))

        # Table for holding File name and number
        doc_para3 = doc_tem.add_paragraph()
        doc_para3.add_run("Data Used").bold = True
        file_table = doc_tem.add_table(rows=1, cols=2)
        file_table.style = "Table Grid"
        first_cell = file_table.rows[0].cells
        first_cell[0].text = self.file_name_entry.get()
        first_cell[1].text = self.file_number_entry.get()

        doc_tem.save("Defect Template.docx")
        messagebox.showinfo("!nfo", "Your file is saved with name Defect Template")

        defect_info()
        self.close_wndw()

    def write_in_defect_file(self):
        old_dir_defect = os.getcwd()
        list_of_images = os.listdir()
        current_dir = old_dir_defect.rstrip("\\Hands-on/Images")
        os.chdir(current_dir)
        try:
            os.mkdir("Defect")
        except: pass
        os.chdir(os.path.join(current_dir,"Defect"))
        img_listt = list()
        doc = docx.Document()
        for x in list_of_images:
            y = os.path.join(old_dir_defect, x)
            img_listt.append(y)
        # Get list of image to save while we pop out the last image to store it below error message in docx
        if len(img_listt) != 0:
            last_pic = img_listt.pop()
            for x in img_listt:
                doc.add_picture(x, width=Cm(10))
            doc_para = doc.add_paragraph()
            doc_para.add_run("ERROR SCREEN and ERROR MESSAGE").bold = True
            doc.add_paragraph(self.actual_result.get(1.0, END))
            doc.add_picture(last_pic, width=Cm(20))
            name = current_dir.split()
            *_, doc_name = name
            doc.save("BUG_FILE_" + doc_name + ".docx")
            messagebox.showinfo("Done!", "Your file is saved with name BUG_FILE_{}".format(doc_name))
            self.defect_template()
        else:
            messagebox.showerror("Error!", "No Screenshot available")
            self.close_wndw()
